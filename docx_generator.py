"""
DOCX Resume & Cover Letter Generator Module — NY DataMind
100% Native Word DOCX Generator using python-docx.
"""

import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any

COLOR_ACCENTS = {
    "Navy Blue": RGBColor(30, 58, 138),
    "Emerald Teal": RGBColor(13, 148, 136),
    "Royal Purple": RGBColor(88, 28, 135),
    "Slate Charcoal": RGBColor(51, 65, 85),
    "Crimson Red": RGBColor(185, 28, 28),
    "Midnight Black": RGBColor(15, 23, 42)
}

def generate_docx(template_key: str, d: Dict[str, Any], custom_color_name="Navy Blue") -> bytes:
    """Generates ultra-professional DOCX binary bytes for candidate profile data."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    accent = COLOR_ACCENTS.get(custom_color_name, RGBColor(30, 58, 138))

    def add_heading(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(title_text.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = accent

    # HEADER SECTION
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(2)
    r_name = p_name.add_run(d.get("full_name", "Alex Morgan"))
    r_name.font.name = 'Calibri'
    r_name.font.size = Pt(22)
    r_name.font.bold = True
    r_name.font.color.rgb = accent

    p_role = doc.add_paragraph()
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_role.paragraph_format.space_after = Pt(4)
    r_role = p_role.add_run(d.get("target_role", "Senior Software Engineer"))
    r_role.font.name = 'Calibri'
    r_role.font.size = Pt(12)
    r_role.font.bold = True
    r_role.font.color.rgb = RGBColor(71, 85, 105)

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(10)
    contact_str = f"{d.get('email', '')} | {d.get('phone', '')} | {d.get('location', '')} | {d.get('linkedin', '')}"
    r_contact = p_contact.add_run(contact_str)
    r_contact.font.name = 'Calibri'
    r_contact.font.size = Pt(9.5)
    r_contact.font.color.rgb = RGBColor(100, 116, 139)

    # 1. SUMMARY
    summary_text = d.get("summary", "")
    if summary_text:
        add_heading("Professional Summary")
        p_sum = doc.add_paragraph()
        p_sum.paragraph_format.space_after = Pt(6)
        r_sum = p_sum.add_run(summary_text)
        r_sum.font.name = 'Calibri'
        r_sum.font.size = Pt(9.5)

    # 2. SKILLS
    skills_text = d.get("skills", "")
    if skills_text:
        add_heading("Technical & Core Skills")
        p_sk = doc.add_paragraph()
        p_sk.paragraph_format.space_after = Pt(6)
        r_sk = p_sk.add_run(f"Core Proficiencies: {skills_text}")
        r_sk.font.name = 'Calibri'
        r_sk.font.size = Pt(9.5)

    # 3. EXPERIENCE
    experiences = d.get("experience", [])
    if experiences:
        add_heading("Professional Experience")
        for exp in experiences:
            p_exp = doc.add_paragraph()
            p_exp.paragraph_format.space_before = Pt(4)
            p_exp.paragraph_format.space_after = Pt(2)
            
            r_title = p_exp.add_run(f"{exp.get('title', '')} - {exp.get('company', '')}")
            r_title.font.name = 'Calibri'
            r_title.font.size = Pt(10.5)
            r_title.font.bold = True
            
            r_dates = p_exp.add_run(f"  ({exp.get('dates', '')} | {exp.get('location', '')})")
            r_dates.font.name = 'Calibri'
            r_dates.font.size = Pt(9.5)
            r_dates.font.italic = True
            r_dates.font.color.rgb = RGBColor(100, 116, 139)

            for bullet in exp.get("bullets", []):
                if bullet.strip():
                    p_b = doc.add_paragraph(style='List Bullet')
                    p_b.paragraph_format.space_after = Pt(2)
                    r_b = p_b.add_run(bullet.strip())
                    r_b.font.name = 'Calibri'
                    r_b.font.size = Pt(9.5)

    # 4. PROJECTS
    projects = d.get("projects", [])
    if projects:
        add_heading("Key Engineering Projects")
        for proj in projects:
            p_pr = doc.add_paragraph()
            p_pr.paragraph_format.space_after = Pt(2)
            r_pr = p_pr.add_run(f"{proj.get('name', '')} [{proj.get('tech', '')}]")
            r_pr.font.name = 'Calibri'
            r_pr.font.size = Pt(10)
            r_pr.font.bold = True

            p_desc = doc.add_paragraph()
            p_desc.paragraph_format.space_after = Pt(4)
            r_desc = p_desc.add_run(proj.get('description', ''))
            r_desc.font.name = 'Calibri'
            r_desc.font.size = Pt(9.5)

    # 5. EDUCATION
    education = d.get("education", [])
    if education:
        add_heading("Education & Credentials")
        for edu in education:
            p_ed = doc.add_paragraph()
            p_ed.paragraph_format.space_after = Pt(4)
            r_ed = p_ed.add_run(f"{edu.get('degree', '')}, {edu.get('institution', '')} ({edu.get('dates', '')})")
            r_ed.font.name = 'Calibri'
            r_ed.font.size = Pt(10)
            r_ed.font.bold = True

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def generate_cover_letter_docx(d: Dict[str, Any], custom_color_name="Navy Blue") -> bytes:
    """Generates matching Cover Letter DOCX binary bytes."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    accent = COLOR_ACCENTS.get(custom_color_name, RGBColor(30, 58, 138))

    name = d.get("full_name", "Alex Morgan")
    role = d.get("target_role", "Software Engineer")
    email = d.get("email", "alex@example.com")
    phone = d.get("phone", "+1 555-0199")
    location = d.get("location", "San Francisco, CA")
    skills = d.get("skills", "Python, SQL, React, AWS")

    # HEADER
    p_name = doc.add_paragraph()
    r_name = p_name.add_run(name)
    r_name.font.name = 'Calibri'
    r_name.font.size = Pt(22)
    r_name.font.bold = True
    r_name.font.color.rgb = accent

    p_role = doc.add_paragraph()
    r_role = p_role.add_run(f"Target Role: {role}")
    r_role.font.name = 'Calibri'
    r_role.font.size = Pt(11)
    r_role.font.bold = True
    r_role.font.color.rgb = RGBColor(71, 85, 105)

    p_contact = doc.add_paragraph()
    p_contact.paragraph_format.space_after = Pt(14)
    r_c = p_contact.add_run(f"Contact: {email} | {phone} | {location}")
    r_c.font.name = 'Calibri'
    r_c.font.size = Pt(9.5)
    r_c.font.color.rgb = RGBColor(100, 116, 139)

    body_paragraphs = [
        "Dear Hiring Manager,",
        f"I am writing to express my strong interest in the {role} position. With a proven track record in technical innovation, system design, and delivering scalable solutions, I am excited about the opportunity to contribute to your team.",
        f"Throughout my career, I have specialized in leveraging core technologies including {skills}. In my previous engineering roles, I have spearheaded critical projects, optimized operational pipelines, and delivered measurable business value.",
        "What drives me is solving complex technical challenges and collaborating with cross-functional teams to build impactful products. I am dedicated to continuous learning, data-driven execution, and modern engineering standards.",
        "I am confident that my technical expertise, problem-solving mindset, and passion for quality make me a strong candidate. I look forward to discussing how my background aligns with your team's goals.",
        f"Thank you for your time and consideration.\n\nSincerely,\n{name}\n{email} | {phone}"
    ]

    for p in body_paragraphs:
        p_para = doc.add_paragraph()
        p_para.paragraph_format.space_after = Pt(8)
        r_p = p_para.add_run(p)
        r_p.font.name = 'Calibri'
        r_p.font.size = Pt(10.5)

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()
